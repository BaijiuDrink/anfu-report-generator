import os
import re
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement, parse_xml
from PIL import Image

_XML_INVALID_CHARS = re.compile(
    "[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufdd0-\ufdef\ufffe-\uffff]"
)


def _sanitize(text):
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    return _XML_INVALID_CHARS.sub("", text)


RISK_COLORS = {
    "严重": RGBColor(139, 0, 0),
    "高危": RGBColor(220, 20, 60),
    "中危": RGBColor(255, 165, 0),
    "低危": RGBColor(0, 100, 0),
    "信息": RGBColor(0, 0, 200),
}

FONT_NAME = "微软雅黑"
CHECKED_BOX = "\u2611"
UNCHECKED_BOX = "\u2610"
SEPARATOR = "\u2501" * 58

ZONE_COLORS = {
    "互联网": RGBColor(0x00, 0x70, 0xC0),
    "内网": RGBColor(0x00, 0x80, 0x40),
}


def _ordered_zones(grouped):
    standard_zones = ["互联网", "内网"]
    return standard_zones + sorted(
        zone for zone in grouped if zone not in standard_zones
    )


def _add_run(paragraph, text, bold=False, size=None, color=None, font_name=FONT_NAME):
    run = paragraph.add_run(_sanitize(text))
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
    rPr.insert(0, rFonts)
    return run


def _add_simple_paragraph(
    doc,
    text,
    bold=False,
    size=11,
    color=None,
    alignment=None,
    space_before=0,
    space_after=4,
    first_line_indent=None,
):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    _add_run(p, text, bold=bold, size=size, color=color)
    return p


def _add_label_row(doc, label, value, label_size=11, value_size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0.5)
    _add_run(
        p, f"{label}：", bold=True, size=label_size, color=RGBColor(0x2F, 0x54, 0x96)
    )
    _add_run(p, value, bold=False, size=value_size)
    return p


def _add_risk_level_row(doc, risk_level):
    levels = ["高危", "中危", "低危"]
    risk_color = RISK_COLORS.get(risk_level, RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)

    _add_run(p, "风险等级：", bold=True, size=11, color=RGBColor(0x2F, 0x54, 0x96))

    for i, level in enumerate(levels):
        box = CHECKED_BOX if level == risk_level else UNCHECKED_BOX
        _add_run(
            p,
            f"  {box}{level}",
            bold=(level == risk_level),
            size=10.5,
            color=risk_color if level == risk_level else None,
        )


def _add_zone_row(doc, zone):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    _add_run(p, "网络区域：", bold=True, size=11, color=RGBColor(0x2F, 0x54, 0x96))
    zone_color = ZONE_COLORS.get(zone, RGBColor(0, 0, 0))
    _add_run(p, zone, bold=True, size=10.5, color=zone_color)


def _add_section_heading(doc, text, level=2):
    if level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, text, bold=True, size=12, color=RGBColor(0x2F, 0x54, 0x96))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = Cm(0.5)
        _add_run(p, text, bold=True, size=11, color=RGBColor(0x33, 0x33, 0x33))
        return p


def _add_body_text(doc, text, indent=1.0, size=10.5):
    if not text:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.first_line_indent = Cm(indent)
        _add_run(p, "（无）", bold=False, size=size, color=RGBColor(150, 150, 150))
        return p

    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.first_line_indent = Cm(indent)
        _add_run(p, line, bold=False, size=size)


def _add_image_to_doc(doc, image_path, max_width_inches=5.5):
    if not image_path or not os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, "[ 无截图 ]", bold=False, size=10, color=RGBColor(180, 180, 180))
        return

    try:
        with Image.open(image_path) as img:
            img_w, img_h = img.size
        max_width_px = max_width_inches * 96
        if img_w > max_width_px:
            ratio = max_width_px / img_w
            width = Inches(max_width_inches)
            height = Inches(img_h * ratio / 96)
        else:
            width = Inches(img_w / 96)
            height = Inches(img_h / 96)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(image_path, width=width, height=height)
    except Exception as e:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, f"[ 截图加载失败: {e} ]", size=9, color=RGBColor(200, 0, 0))


def _add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, SEPARATOR, size=8, color=RGBColor(180, 180, 180))


class ReportBuilder:
    def __init__(self, project_name="渗透测试报告"):
        self.document = Document()
        self.project_name = _sanitize(project_name)
        self._disable_image_compression()
        self._setup_page()
        self._add_cover()

    def _disable_image_compression(self):
        settings = self.document.settings.element
        if settings.find(qn("w:doNotCompressPictures")) is None:
            settings.append(OxmlElement("w:doNotCompressPictures"))

    def _setup_page(self):
        section = self.document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    def _add_cover(self):
        for _ in range(4):
            self.document.add_paragraph()

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        _add_run(p, self.project_name, bold=True, size=22)

        p2 = self.document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        _add_run(
            p2, "渗透测试报告", bold=False, size=16, color=RGBColor(0x2F, 0x54, 0x96)
        )

        for _ in range(2):
            self.document.add_paragraph()

        now = datetime.datetime.now()
        info_lines = [
            f'报告编号：PT-{now.strftime("%Y%m%d")}-001',
            f'测试日期：{now.strftime("%Y年%m月%d日")}',
            f'报告日期：{now.strftime("%Y年%m月%d日")}',
            "密级：内部",
        ]
        for line in info_lines:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            _add_run(p, line, size=11, color=RGBColor(80, 80, 80))

        self.document.add_page_break()

    def _count_by_zone(self, findings):
        counts = {}
        for f in findings:
            zone = f.get("network_zone", "互联网")
            level = f.get("risk_level", "中危")
            if zone not in counts:
                counts[zone] = {
                    "total": 0,
                    "严重": 0,
                    "高危": 0,
                    "中危": 0,
                    "低危": 0,
                    "信息": 0,
                }
            counts[zone]["total"] += 1
            if level in counts[zone]:
                counts[zone][level] += 1
        return counts

    def add_summary_section(self, findings):
        total = len(findings)
        if total == 0:
            return

        zone_counts = self._count_by_zone(findings)

        _add_simple_paragraph(
            self.document,
            "漏洞统计概要",
            bold=True,
            size=16,
            color=RGBColor(0x2F, 0x54, 0x96),
            space_before=0,
            space_after=10,
        )

        sum_p = self.document.add_paragraph()
        sum_p.paragraph_format.space_after = Pt(4)
        _add_run(sum_p, f"本次渗透测试共发现 {total} 个安全漏洞。", size=11)

        for zone in _ordered_zones(zone_counts):
            if zone not in zone_counts:
                continue
            zc = zone_counts[zone]
            zone_color = ZONE_COLORS.get(zone, RGBColor(0, 0, 0))
            z_p = self.document.add_paragraph()
            z_p.paragraph_format.space_after = Pt(3)
            _add_run(z_p, f"▎{zone} ", bold=True, size=11, color=zone_color)
            _add_run(z_p, f'共 {zc["total"]} 个漏洞：', size=11)
            parts = []
            for level in ["严重", "高危", "中危", "低危", "信息"]:
                if zc[level] > 0:
                    color = RISK_COLORS.get(level, RGBColor(0, 0, 0))
                    parts.append((f"{level}{zc[level]}个", color))
            for i, (text, color) in enumerate(parts):
                if i > 0:
                    _add_run(z_p, "、", size=11)
                _add_run(z_p, text, bold=True, size=11, color=color)

        self.document.add_page_break()

    def add_zone_heading(self, zone, index=1):
        zone_color = ZONE_COLORS.get(zone, RGBColor(0x2F, 0x54, 0x96))
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        _add_run(p, f'{"=" * 50}', size=10, color=RGBColor(200, 200, 200))
        p2 = self.document.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        _add_run(
            p2,
            f"第{index}部分：{zone}区域漏洞详情",
            bold=True,
            size=15,
            color=zone_color,
        )

    def add_finding_detail(self, finding, index=1):
        name = finding.get("name", "未知漏洞")
        url = finding.get("url", "")
        host_ip = finding.get("host_ip", "")
        risk_level = finding.get("risk_level", "中危")
        network_zone = finding.get("network_zone", "互联网")
        description = finding.get("description", "")

        verify_steps = finding.get("verify_steps", "")
        verify_result = finding.get("verify_result", "")

        fix_suggestion = finding.get("fix_suggestion", "")
        fix_priority = finding.get("fix_priority", "高")
        fix_verify = finding.get("fix_verify", "")

        _add_section_heading(self.document, f"{index}. 漏洞详情", level=2)

        _add_label_row(self.document, "漏洞名称", name)

        _add_zone_row(self.document, network_zone)

        display_url = url or host_ip or "—"
        _add_label_row(self.document, "漏洞地址", display_url)

        if host_ip and host_ip != url:
            _add_label_row(self.document, "主机IP", host_ip)

        _add_risk_level_row(self.document, risk_level)

        _add_section_heading(self.document, "漏洞描述", level=3)
        description = (
            description
            or "详细说明漏洞成因、触发条件。示例：被测系统登录接口未对输入参数进行过滤，攻击者可通过注入恶意SQL语句，绕过身份认证获取后台管理权限，进而访问敏感数据。"
        )
        desc_text = f"详细说明漏洞成因、触发条件，示例：{description}"
        _add_body_text(self.document, desc_text)

        _add_section_heading(self.document, "漏洞验证", level=3)
        self._add_verification_section(verify_steps, verify_result)

        _add_section_heading(self.document, "修复建议", level=3)

        fix_suggestion = fix_suggestion or "（技术措施，待补充）"
        priorities = ["紧急", "高", "中", "低"]
        priority_labels = {
            "紧急": "紧急（24h内）",
            "高": "高（3个工作日）",
            "中": "中（1周内）",
            "低": "低（下版本迭代）",
        }

        _add_body_text(
            self.document,
            f"技术措施（可落地，如：对输入参数进行正则过滤，禁用危险SQL函数，使用预编译语句）：{fix_suggestion}",
        )

        p_priority = self.document.add_paragraph()
        p_priority.paragraph_format.space_before = Pt(2)
        p_priority.paragraph_format.space_after = Pt(2)
        p_priority.paragraph_format.first_line_indent = Cm(1.0)
        _add_run(p_priority, "整改优先级（", bold=False, size=10.5)
        for i, p in enumerate(priorities):
            box = CHECKED_BOX if p == fix_priority else UNCHECKED_BOX
            _add_run(p_priority, f"{box}{priority_labels.get(p, p)}", size=10.5)
            if i < len(priorities) - 1:
                _add_run(p_priority, "  ", size=10.5)
        _add_run(p_priority, "）", bold=False, size=10.5)

        fix_verify = (
            fix_verify
            or "（整改验证方法，如：使用原测试工具+手工验证，无漏洞触发即为整改完成）"
        )
        _add_body_text(
            self.document,
            f"整改验证方法（如：使用原测试工具+手工验证，无漏洞触发即为整改完成）：{fix_verify}",
        )

        _add_separator(self.document)

    def _add_verification_section(self, verify_steps, verify_result):
        verify_steps = verify_steps or "（含Payload、操作流程，待补充）"
        verify_result = verify_result or "（验证结果，待补充）"

        import re

        img_pattern = re.compile(
            r"\[截图:\s*([^\]]+\.(?:png|jpg|jpeg|gif|bmp))\]", re.IGNORECASE
        )

        if img_pattern.search(verify_steps):
            parts = img_pattern.split(verify_steps)
            for i, part in enumerate(parts):
                if not part:
                    continue
                if i % 2 == 0:
                    text = part.strip()
                    if text:
                        _add_body_text(self.document, text)
                else:
                    img_path = part.strip()
                    base_dir = (
                        os.path.dirname(os.path.abspath(__file__))
                        if not os.path.isabs(img_path)
                        else ""
                    )
                    full_path = (
                        os.path.join(base_dir, img_path) if base_dir else img_path
                    )
                    _add_image_to_doc(self.document, full_path)

            _add_body_text(
                self.document, "验证步骤（含Payload、操作流程）：见上方截图及描述"
            )
        else:
            _add_body_text(
                self.document, f"验证步骤（含Payload、操作流程）：{verify_steps}"
            )

        _add_body_text(
            self.document,
            f"验证结果（如：成功获取管理员账号密码，截图见附件）：{verify_result}",
        )

    def add_findings_section(self, findings):
        if not findings:
            _add_simple_paragraph(
                self.document,
                "未录入漏洞信息",
                size=12,
                color=RGBColor(150, 150, 150),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            return

        grouped = {"互联网": [], "内网": []}
        for f in findings:
            zone = f.get("network_zone", "互联网")
            if zone not in grouped:
                grouped[zone] = []
            grouped[zone].append(f)

        part_num = 1
        global_idx = 1
        for zone in _ordered_zones(grouped):
            zone_findings = grouped.get(zone, [])
            if not zone_findings:
                continue
            self.add_zone_heading(zone, part_num)
            part_num += 1
            for f in zone_findings:
                self.add_finding_detail(f, global_idx)
                global_idx += 1

    def save(self, output_path):
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        self.document.save(output_path)
        return output_path
